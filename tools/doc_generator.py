"""Document generator for creating professional Microsoft Word documents."""

import logging
from datetime import datetime
from pathlib import Path

from docx import Document

from config import config
from models.schemas import DocumentSection

logger = logging.getLogger(__name__)

DOCUMENT_TITLE = "Generated Document"


def create_document(sections: list[DocumentSection]) -> str:
    """Generate a Word document from structured sections and return the file path."""
    if not sections:
        logger.warning("create_document called with empty sections")
        return ""

    doc = Document()

    doc.add_heading(DOCUMENT_TITLE, level=1)

    for section in sections:
        doc.add_heading(section.title, level=2)
        doc.add_paragraph(section.content)

    output_dir = Path(config.OUTPUT_DIRECTORY)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"document_{timestamp}.docx"
    filepath = output_dir / filename

    doc.save(str(filepath))
    logger.info("Document saved to %s", filepath)
    return str(filepath)
